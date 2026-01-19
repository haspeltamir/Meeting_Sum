import os
from datetime import datetime
from docx import Document
from docx.shared import Pt

def generate_docx(job_data: dict, output_path: str):
    """
    Generates a Word document from the job analysis data.
    """
    document = Document()

    # Metadata
    job_id = job_data.get("job_id", "Unknown")
    created_at = job_data.get("created_at", datetime.now())
    if isinstance(created_at, str):
         # basic parsing if it's a string, though usually datetime object from mongo
         try:
             created_at = datetime.fromisoformat(created_at)
         except ValueError:
             created_at = datetime.now()
    
    date_str = created_at.strftime("%Y-%m-%d %H:%M")

    # Heading 1: Title
    document.add_heading('Meeting Executive Summary', 0)
    
    # Metadata Sub-header
    p = document.add_paragraph()
    p.add_run(f"Date: {date_str}").italic = True
    p.add_run(f"\nJob ID: {job_id}").italic = True

    # Analysis Data
    analysis = job_data.get("analysis", {})
    if not analysis:
        document.add_paragraph("No analysis data available.")
        document.save(output_path)
        return

    # Heading 2: Summary
    document.add_heading('Summary', level=1)
    document.add_paragraph(analysis.get("summary", "N/A"))

    # Heading 2: Participants
    document.add_heading('Participants', level=1)
    participants = analysis.get("participants", [])
    if participants:
        for person in participants:
            document.add_paragraph(person, style='List Bullet')
    else:
        document.add_paragraph("None listed.")

    # Heading 2: Decisions
    document.add_heading('Decisions', level=1)
    decisions = analysis.get("decisions", [])
    if decisions:
        for decision in decisions:
            document.add_paragraph(decision, style='List Bullet')
    else:
        document.add_paragraph("None listed.")

    # Heading 2: Action Items
    document.add_heading('Action Items', level=1)
    action_items = analysis.get("action_items", [])
    
    if action_items:
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Task'
        hdr_cells[1].text = 'Owner'
        hdr_cells[2].text = 'Due Date'

        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("task", "")
            row_cells[1].text = item.get("owner") or ""
            row_cells[2].text = item.get("due_date") or ""
    else:
        document.add_paragraph("None listed.")

    # Heading 2: Full Transcript
    document.add_heading('Full Transcript', level=1)
    document.add_paragraph(job_data.get("transcript", "N/A"))

    # Save
    document.save(output_path)
